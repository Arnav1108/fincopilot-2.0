import csv
import json
import os
import time
import uuid
from datetime import datetime

import openai
import pypdf
import redis
import structlog
import tiktoken
from bs4 import BeautifulSoup
from celery.exceptions import Retry
from docx import Document as DocxDocument
from docx.oxml.ns import qn as _oxml_qn
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker

from app.celery_app import celery_app
from app.config import settings
from app.models.document import Document, DocumentChunk, DocumentStatus

_SYNC_DATABASE_URL = settings.DATABASE_URL.replace(
    "postgresql+asyncpg://", "postgresql+psycopg2://"
)
_engine = create_engine(_SYNC_DATABASE_URL, pool_pre_ping=True)
_SessionLocal = sessionmaker(bind=_engine)
_redis_client = redis.Redis.from_url(settings.REDIS_URL, decode_responses=True)
_task_log = structlog.get_logger(__name__)


# ── Status publisher ──────────────────────────────────────────────────────────

def _publish_status(
    document_id: str,
    status: str,
    chunk_count: int | None = None,
    error_message: str | None = None,
) -> None:
    channel = f"document.{document_id}.status"
    try:
        _redis_client.publish(
            channel,
            json.dumps({"status": status, "chunk_count": chunk_count, "error_message": error_message}),
        )
        _task_log.debug("redis_status_published", document_id=document_id, status=status)
    except Exception as exc:
        _task_log.warning("redis_publish_failed", document_id=document_id, error=str(exc))


# ── Per-format parsers ────────────────────────────────────────────────────────
# Each returns list[tuple[page_num, text]] to fit the existing chunking logic.

def _parse_pdf(file_path: str) -> list[tuple[int, str]]:
    reader = pypdf.PdfReader(file_path)
    return [(i, page.extract_text() or "") for i, page in enumerate(reader.pages)]


def _parse_docx(file_path: str) -> list[tuple[int, str]]:
    doc = DocxDocument(file_path)
    segments: list[str] = []

    # Paragraphs — preserve heading level as Markdown prefix
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name
        if style.startswith("Heading 1"):
            segments.append(f"# {text}")
        elif style.startswith("Heading 2"):
            segments.append(f"## {text}")
        elif style.startswith("Heading"):
            segments.append(f"### {text}")
        else:
            segments.append(text)

    # Tables — each cell separated by " | ", rows by newline
    for table in doc.tables:
        rows = [
            " | ".join(cell.text.strip() for cell in row.cells)
            for row in table.rows
        ]
        segments.append("\n".join(rows))

    # Text boxes stored in the body XML
    for txbx in doc.element.body.iter(_oxml_qn("w:txbxContent")):
        text = " ".join(t.text for t in txbx.iter(_oxml_qn("w:t")) if t.text).strip()
        if text:
            segments.append(text)

    # Group segments into pages of ~30 each so page_numbers in metadata are useful
    PAGE_SIZE = 30
    pages: list[tuple[int, str]] = []
    for i in range(0, max(len(segments), 1), PAGE_SIZE):
        group = segments[i : i + PAGE_SIZE]
        pages.append((i // PAGE_SIZE, "\n".join(group)))

    return pages or [(0, "")]


def _parse_csv(file_path: str) -> list[tuple[int, str]]:
    ROWS_PER_PAGE = 10
    pages: list[tuple[int, str]] = []

    with open(file_path, newline="", encoding="utf-8-sig") as fh:
        reader = csv.DictReader(fh)
        headers = reader.fieldnames or []
        header_line = " | ".join(headers)
        page_num = 0
        batch: list[str] = []

        for row in reader:
            batch.append(" | ".join(str(row.get(h, "")) for h in headers))
            if len(batch) >= ROWS_PER_PAGE:
                pages.append((page_num, header_line + "\n" + "\n".join(batch)))
                page_num += 1
                batch = []

        if batch:
            pages.append((page_num, header_line + "\n" + "\n".join(batch)))

    return pages or [(0, "")]


def _parse_txt(file_path: str) -> list[tuple[int, str]]:
    with open(file_path, encoding="utf-8", errors="replace") as fh:
        return [(0, fh.read())]


def _parse_html(file_path: str) -> list[tuple[int, str]]:
    with open(file_path, encoding="utf-8", errors="replace") as fh:
        raw = fh.read()
    soup = BeautifulSoup(raw, "html.parser")
    for tag in soup(["script", "style"]):
        tag.decompose()
    parts: list[str] = []
    for el in soup.find_all(["p", "div", "h1", "h2", "h3", "h4", "h5", "h6", "table"]):
        text = el.get_text(separator=" ", strip=True)
        if text:
            parts.append(text)
    return [(0, "\n".join(parts))]


def _extract_pages(file_path: str, file_type: str) -> list[tuple[int, str]]:
    dispatch = {
        "pdf":  _parse_pdf,
        "docx": _parse_docx,
        "csv":  _parse_csv,
        "txt":  _parse_txt,
        "html": _parse_html,
    }
    parser = dispatch.get(file_type.lower())
    if parser is None:
        raise ValueError(f"Unsupported file_type: {file_type!r}")
    return parser(file_path)


# ── Chunker ───────────────────────────────────────────────────────────────────

def _build_chunks(full_text: str, pages: list[tuple[int, str]]) -> list[dict]:
    enc = tiktoken.get_encoding("cl100k_base")
    all_tokens = enc.encode(full_text)

    # Cumulative char offset at each token boundary
    token_char_offsets = [0]
    running = 0
    for token_id in all_tokens:
        token_bytes = enc.decode_single_token_bytes(token_id)
        try:
            char_len = len(token_bytes.decode("utf-8"))
        except UnicodeDecodeError:
            char_len = len(token_bytes)
        running += char_len
        token_char_offsets.append(running)

    # Character range per page in the concatenated full_text ("\n"-joined)
    page_boundaries: list[tuple[int, int, int]] = []
    cumulative = 0
    for page_num, page_text in pages:
        start = cumulative
        end = cumulative + len(page_text)
        page_boundaries.append((start, end, page_num))
        cumulative = end + 1  # +1 for the "\n" separator between pages

    MAX_TOKENS = 800
    OVERLAP = 100
    chunks: list[dict] = []
    start = 0
    total = len(all_tokens)

    while start < total:
        end = min(start + MAX_TOKENS, total)
        chunk_text = enc.decode(all_tokens[start:end])
        chunk_start_char = token_char_offsets[start]
        chunk_end_char = token_char_offsets[end]

        page_nums = sorted(
            {
                pn
                for (ps, pe, pn) in page_boundaries
                if ps < chunk_end_char and pe > chunk_start_char
            }
        )

        chunks.append(
            {
                "chunk_index": len(chunks),
                "content": chunk_text,
                "page_numbers": page_nums,
            }
        )

        if end == total:
            break
        start = end - OVERLAP

    return chunks


# ── Task ──────────────────────────────────────────────────────────────────────

@celery_app.task(bind=True, max_retries=3, name="app.tasks.ingestion.ingest_document")
def ingest_document(
    self,
    document_id: str,
    file_path: str,
    file_type: str,
    conversation_id: str,
) -> None:
    log = structlog.get_logger(__name__).bind(
        document_id=document_id, file_type=file_type
    )
    log.info("task_started")

    start_time = time.monotonic()
    db = _SessionLocal()

    try:
        # 1. Fetch document and set processing
        doc = db.get(Document, uuid.UUID(document_id))
        doc.status = DocumentStatus.processing
        doc.updated_at = datetime.utcnow()
        db.commit()
        _publish_status(document_id, "processing")

        # 2. Parse into (page_num, text) tuples — format-aware
        pages = _extract_pages(file_path, file_type)
        full_text = "\n".join(text for _, text in pages)
        log.info(
            "extraction_complete",
            page_count=len(pages),
            char_count=len(full_text),
        )

        # 3. Guard: not enough extractable text
        if len(full_text.strip()) < 50:
            doc.status = DocumentStatus.failed
            doc.error_message = "no extractable text"
            doc.updated_at = datetime.utcnow()
            db.commit()
            _publish_status(document_id, "failed", error_message="no extractable text")
            return

        # 4. Chunk with sliding window (800 tokens, 100 overlap)
        chunk_dicts = _build_chunks(full_text, pages)
        log.info("chunking_complete", chunk_count=len(chunk_dicts))

        # 5. Embed all chunks in a single API call
        client = openai.OpenAI(api_key=settings.OPENAI_API_KEY)
        try:
            response = client.embeddings.create(
                model="text-embedding-3-small",
                input=[c["content"] for c in chunk_dicts],
            )
        except (openai.APIError, openai.RateLimitError) as exc:
            db.rollback()
            if self.request.retries >= self.max_retries:
                doc = db.get(Document, uuid.UUID(document_id))
                doc.status = DocumentStatus.failed
                doc.error_message = str(exc)
                doc.updated_at = datetime.utcnow()
                db.commit()
                _publish_status(document_id, "failed", error_message=str(exc))
                log.error("task_failed", error=str(exc))
                return
            log.warning(
                "task_failed_retrying",
                error=str(exc),
                retry_number=self.request.retries,
            )
            raise self.retry(exc=exc, countdown=30 * (2 ** self.request.retries))

        log.info("embedding_complete")

        # 6. Build DocumentChunk objects
        conv_uuid = uuid.UUID(conversation_id)
        chunks = [
            DocumentChunk(
                document_id=doc.id,
                user_id=doc.user_id,
                conversation_id=conv_uuid,
                chunk_index=c["chunk_index"],
                content=c["content"],
                embedding=response.data[i].embedding,
                chunk_metadata={
                    "page_numbers": c["page_numbers"],
                    "chunk_index": c["chunk_index"],
                    "source_filename": doc.filename,
                    "file_type": file_type,
                },
            )
            for i, c in enumerate(chunk_dicts)
        ]

        # 7. Single bulk insert
        db.add_all(chunks)
        db.flush()

        # 8. Mark document ready
        doc.chunk_count = len(chunks)
        doc.status = DocumentStatus.ready
        doc.updated_at = datetime.utcnow()
        db.commit()
        _publish_status(document_id, "ready", chunk_count=len(chunks))

        elapsed = time.monotonic() - start_time
        log.info(
            "task_succeeded",
            chunk_count=len(chunks),
            elapsed_seconds=round(elapsed, 2),
        )

    except Retry:
        raise

    except Exception as exc:
        db.rollback()
        log.error("task_failed", error=str(exc))
        try:
            doc = db.get(Document, uuid.UUID(document_id))
            if doc is not None:
                doc.status = DocumentStatus.failed
                doc.error_message = str(exc)
                doc.updated_at = datetime.utcnow()
                db.commit()
                _publish_status(document_id, "failed", error_message=str(exc))
        except Exception:
            pass

    finally:
        db.close()
        if os.path.exists(file_path):
            os.remove(file_path)
